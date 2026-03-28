"""
generar_memoria_word.py
════════════════════════════════════════════════════════════════════════
Generador de Memoria de Cálculo Word — NOM-081-SEMARNAT-1994
100 % Python, sin dependencias de Node.js.

Uso:
    from generar_memoria_word import generar_word
    docx_bytes = generar_word(resultado_dict)   # → bytes listos para descarga
════════════════════════════════════════════════════════════════════════
"""

import io
import math
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy
except ImportError as e:
    raise ImportError(
        f"Librería faltante: {e}\n"
        "Instala con:  pip install python-docx"
    )

# ── Paleta de colores (RGB) ────────────────────────────────────────────────────
C_GREEN_TITLE  = RGBColor(0xE2, 0xEF, 0xD9)   # verde título principal
C_GREEN_HEADER = RGBColor(0xC6, 0xE0, 0xB4)   # verde encabezados tabla
C_BLUE_HEADER  = RGBColor(0xD9, 0xE1, 0xF2)   # azul claro fondo
C_GREY_ROW     = RGBColor(0xF2, 0xF2, 0xF2)   # filas alternas
C_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK_GREEN   = RGBColor(0x37, 0x56, 0x23)   # texto verde oscuro
C_RED_FAIL     = RGBColor(0xFF, 0xCC, 0xCC)   # fondo rojo (excede)
C_GREEN_PASS   = RGBColor(0xC6, 0xEF, 0xCE)   # fondo verde (cumple)
C_RESULT_TEXT_OK   = RGBColor(0x37, 0x56, 0x23)
C_RESULT_TEXT_FAIL = RGBColor(0xC0, 0x00, 0x00)

FONT_NAME = "Arial Nova Light"
FONT_FALL = "Arial"


# ── Helpers XML (bordes, sombreado) ───────────────────────────────────────────

def _rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _set_cell_bg(cell, color: RGBColor):
    """Aplica color de fondo.
    Orden en tcPr: tcW, gridSpan, vMerge, tcBorders, shd, noWrap, tcMar, vAlign, hideMark
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_hex(color))
    # Preferir insertar después de tcBorders
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is not None:
        tcBorders.addnext(shd)
        return
    # Si no hay tcBorders, insertar antes de noWrap/tcMar/vAlign/hideMark
    for anchor_tag in [qn("w:noWrap"), qn("w:tcMar"), qn("w:vAlign"), qn("w:hideMark")]:
        anchor = tcPr.find(anchor_tag)
        if anchor is not None:
            anchor.addprevious(shd)
            return
    tcPr.append(shd)


def _set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                      size=4, color="000000"):
    """Aplica bordes individuales a una celda."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")
    # OOXML 2006 strict order: top, start, bottom, end
    for side, show in [("top", top), ("start", left),
                       ("bottom", bottom), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        if show:
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    str(size))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    # Insert tcBorders BEFORE shd (shd must come after tcBorders in tcPr)
    shd_el = tcPr.find(qn("w:shd"))
    if shd_el is not None:
        shd_el.addprevious(borders)
    else:
        tcPr.append(borders)


def _set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    """Aplica margenes internos a la celda (twips).
    Schema order: tcBorders -> shd -> noWrap -> tcMar -> vAlign -> hideMark
    tcMar debe insertarse ANTES de vAlign/hideMark.
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(existing)
    mar = OxmlElement("w:tcMar")
    # OOXML 2006 strict order inside tcMar: top, start, bottom, end
    for side, val in [("top", top), ("start", left),
                      ("bottom", bottom), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    # Insert BEFORE vAlign / hideMark to respect schema ordering
    for anchor_tag in [qn("w:vAlign"), qn("w:hideMark")]:
        anchor = tcPr.find(anchor_tag)
        if anchor is not None:
            anchor.addprevious(mar)
            return
    tcPr.append(mar)


def _set_col_width(cell, width_cm: float):
    """Establece ancho de celda en cm."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    twips = int(width_cm * 567)   # 1 cm ≈ 567 twips
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _cell_text(cell, text: str, bold=False, size=10, align="center",
               color: Optional[RGBColor] = None, italic=False):
    """Limpia la celda y escribe texto con formato."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    run = p.add_run(str(text))
    run.bold    = bold
    run.italic  = italic
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

    # Espaciado del párrafo
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)

    # Alineación vertical
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return run


def _style_all_cells(table):
    """Aplica bordes y márgenes a todas las celdas de una tabla."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)
            _set_cell_margins(cell)


# ── Configuración de página ────────────────────────────────────────────────────

def _setup_page(doc):
    """Configura página carta con márgenes de 1.9 cm."""
    from docx.oxml.ns import nsmap
    section = doc.sections[0]
    section.page_width  = Cm(21.59)   # Letter 8.5"
    section.page_height = Cm(27.94)   # Letter 11"
    section.top_margin    = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.left_margin   = Cm(1.9)
    section.right_margin  = Cm(1.9)


def _add_footer(doc, meta):
    """Agrega pie de página con línea divisoria."""
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False

    p    = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Línea superior sobre el footer
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top_el = OxmlElement("w:top")
    top_el.set(qn("w:val"),   "single")
    top_el.set(qn("w:sz"),    "4")
    top_el.set(qn("w:space"), "4")
    top_el.set(qn("w:color"), "375623")
    pBdr.append(top_el)
    # pBdr debe ir antes de shd, tabs, suppressLineNumbers en pPr
    # Insertar antes de cualquier elemento que venga después en el schema
    inserted = False
    for anchor_tag in [qn("w:shd"), qn("w:tabs"), qn("w:jc"),
                       qn("w:rPr"), qn("w:spacing"), qn("w:ind")]:
        anchor = pPr.find(anchor_tag)
        if anchor is not None:
            anchor.addprevious(pBdr)
            inserted = True
            break
    if not inserted:
        pPr.append(pBdr)

    run = p.add_run(
        f"NOM-081-SEMARNAT-1994  |  {meta.get('compania','')}  |  {meta.get('fecha','')}"
    )
    run.font.name  = FONT_FALL
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ── Párrafo con estilo ─────────────────────────────────────────────────────────

def _add_para(doc, text, bold=False, size=10, align="left",
              space_before=0, space_after=4,
              color: Optional[RGBColor] = None, italic=False,
              indent_left=0):
    p = doc.add_paragraph()
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if indent_left:
        pf.left_indent = Pt(indent_left)

    run = p.add_run(text)
    run.bold    = bold
    run.italic  = italic
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _add_formula(doc, text):
    """Párrafo centrado en itálica para fórmulas."""
    return _add_para(doc, text, italic=True, size=10,
                     align="center", space_before=4, space_after=4)


def _add_indent(doc, text, size=10, bold=False):
    """Párrafo con sangría izquierda para explicaciones."""
    return _add_para(doc, text, size=size, bold=bold,
                     indent_left=18, space_before=2, space_after=2)


# ── TABLAS ────────────────────────────────────────────────────────────────────

def _tabla_encabezado(doc, meta):
    """Tabla de identificación del estudio (ficha de datos)."""
    tbl = doc.add_table(rows=4, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [4.0, 10.5, 4.5]
    for i, row in enumerate(tbl.rows):
        for j, cell in enumerate(row.cells):
            _set_col_width(cell, widths[j])
            _set_cell_borders(cell)
            _set_cell_margins(cell)

    # Fila 0 — Compañía
    _cell_text(tbl.cell(0, 0), "Compañía:", bold=True, size=9, align="right")
    _set_cell_bg(tbl.cell(0, 0), C_GREY_ROW)
    _cell_text(tbl.cell(0, 1), meta.get("compania", ""), size=9, align="left")
    _cell_text(tbl.cell(0, 2), "Fecha de muestreo:", bold=True, size=9, align="center")
    _set_cell_bg(tbl.cell(0, 2), C_GREY_ROW)

    # Fila 1 — Ubicación
    _cell_text(tbl.cell(1, 0), "Ubicación:", bold=True, size=9, align="right")
    _set_cell_bg(tbl.cell(1, 0), C_GREY_ROW)
    _cell_text(tbl.cell(1, 1), meta.get("ubicacion", ""), size=9, align="left")
    _cell_text(tbl.cell(1, 2), meta.get("fecha", ""), size=9, align="center")

    # Fila 2 — Evaluadores
    _cell_text(tbl.cell(2, 0), "Evaluadores:", bold=True, size=9, align="right")
    _set_cell_bg(tbl.cell(2, 0), C_GREY_ROW)
    _cell_text(tbl.cell(2, 1), meta.get("evaluadores", ""), size=9, align="left")
    ev_txt = (f"Evaluación: {meta.get('evaluacion','Diurna')}\n"
              f"Inicio: {meta.get('hora_inicio','')}\n"
              f"Final:  {meta.get('hora_final','')}")
    _cell_text(tbl.cell(2, 2), ev_txt, size=9, align="left")
    _set_cell_bg(tbl.cell(2, 2), C_GREY_ROW)

    # Fila 3 — Zona
    _cell_text(tbl.cell(3, 0), "Zona Crítica:", bold=True, size=9, align="right")
    _set_cell_bg(tbl.cell(3, 0), C_GREY_ROW)
    _cell_text(tbl.cell(3, 1), meta.get("zona", ""), size=9, align="left")
    lim = meta.get("limite", 68.0)
    _cell_text(tbl.cell(3, 2), f"Límite permisible: {lim} dB",
               bold=True, size=9, align="center")
    _set_cell_bg(tbl.cell(3, 2), C_GREEN_HEADER)


def _tabla_datos_campo(doc, tipo, periodos, raw_data, stats):
    """Tabla de 35 lecturas + sumatoria para Fuente o Fondo."""
    n_cols = len(periodos) + 1

    # ── Título bloque ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(2)
    r = p.add_run(f"DATOS DE CAMPO — RUIDO DE {tipo}")
    r.bold = True; r.font.name = FONT_NAME; r.font.size = Pt(10)

    tbl = doc.add_table(rows=3 + 35 + 1, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Anchos de columna
    col_w_lec = 2.0
    col_w_per = (17.0 - col_w_lec) / len(periodos)
    for row in tbl.rows:
        _set_col_width(row.cells[0], col_w_lec)
        for j in range(1, n_cols):
            _set_col_width(row.cells[j], col_w_per)

    # ── Fila 0: encabezado "Período" fusionado ──
    r0 = tbl.rows[0]
    _cell_text(r0.cells[0], "N° de\nLectura", bold=True, size=9)
    _set_cell_bg(r0.cells[0], C_GREEN_HEADER)
    _cell_text(r0.cells[1], "Periodo", bold=True, size=9)
    _set_cell_bg(r0.cells[1], C_GREEN_HEADER)
    # Fusionar celdas de períodos
    if len(periodos) > 1:
        r0.cells[1].merge(r0.cells[n_cols - 1])

    # ── Fila 1: letras/números de período ──
    r1 = tbl.rows[1]
    _cell_text(r1.cells[0], "", size=9)
    _set_cell_bg(r1.cells[0], C_GREEN_HEADER)
    for j, p_name in enumerate(periodos):
        _cell_text(r1.cells[j + 1], p_name, bold=True, size=9)
        _set_cell_bg(r1.cells[j + 1], C_GREEN_HEADER)

    # ── Fila 2: encabezado "N° / valores" ──
    r2 = tbl.rows[2]
    _cell_text(r2.cells[0], "N°", bold=True, size=9)
    _set_cell_bg(r2.cells[0], C_GREEN_HEADER)
    for j, p_name in enumerate(periodos):
        _cell_text(r2.cells[j + 1], p_name, bold=True, size=9)
        _set_cell_bg(r2.cells[j + 1], C_GREEN_HEADER)

    # ── 35 filas de lecturas ──
    for i in range(35):
        row = tbl.rows[3 + i]
        fill = C_WHITE if i % 2 == 0 else C_GREY_ROW
        _cell_text(row.cells[0], str(i + 1), size=9)
        _set_cell_bg(row.cells[0], C_GREEN_HEADER)
        for j, p_name in enumerate(periodos):
            val = raw_data[p_name][i]
            _cell_text(row.cells[j + 1], f"{val:.1f}", size=9)
            _set_cell_bg(row.cells[j + 1], fill)

    # ── Fila sumatoria ──
    r_sum = tbl.rows[-1]
    _cell_text(r_sum.cells[0], "Sumatoria", bold=True, size=9)
    _set_cell_bg(r_sum.cells[0], C_GREEN_HEADER)
    for j, p_name in enumerate(periodos):
        _cell_text(r_sum.cells[j + 1], f"{stats[p_name]['suma']:.2f}", bold=True, size=9)
        _set_cell_bg(r_sum.cells[j + 1], C_GREEN_TITLE)

    _style_all_cells(tbl)


def _tabla_resultados_periodo(doc, tipo, periodos, stats, fill_hdr):
    """Tabla de N50, σ, N10, Neq por período."""
    n_cols = len(periodos) + 1

    tbl = doc.add_table(rows=1 + 1 + 4, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w0 = 3.5
    col_wp = (17.0 - col_w0) / len(periodos)

    for row in tbl.rows:
        _set_col_width(row.cells[0], col_w0)
        for j in range(1, n_cols):
            _set_col_width(row.cells[j], col_wp)

    # Título fusionado
    title_row = tbl.rows[0]
    title_row.cells[0].merge(title_row.cells[n_cols - 1])
    _cell_text(title_row.cells[0],
               f"RESULTADOS POR PERÍODO — RUIDO DE {tipo}",
               bold=True, size=9)
    _set_cell_bg(title_row.cells[0], fill_hdr)

    # Encabezado
    hdr = tbl.rows[1]
    _cell_text(hdr.cells[0], "Indicador", bold=True, size=9)
    _set_cell_bg(hdr.cells[0], C_GREEN_HEADER)
    for j, p_name in enumerate(periodos):
        _cell_text(hdr.cells[j + 1], p_name, bold=True, size=9)
        _set_cell_bg(hdr.cells[j + 1], C_GREEN_HEADER)

    indicadores = [
        ("N₅₀ (dB)", "N50"),
        ("σ (dB)",   "sigma"),
        ("N₁₀ (dB)", "N10"),
        ("Neq (dB)", "Neq"),
    ]
    for i, (label, key) in enumerate(indicadores):
        row = tbl.rows[2 + i]
        fill = C_WHITE if i % 2 == 0 else C_GREY_ROW
        _cell_text(row.cells[0], label, bold=True, size=9, align="left")
        _set_cell_bg(row.cells[0], C_GREEN_HEADER)
        for j, p_name in enumerate(periodos):
            _cell_text(row.cells[j + 1], f"{stats[p_name][key]:.2f}", size=9)
            _set_cell_bg(row.cells[j + 1], fill)

    _style_all_cells(tbl)


def _tabla_promedios(doc, prom):
    """Tabla comparativa de promedios Fuente vs Fondo."""
    tbl = doc.add_table(rows=1 + 1 + 4, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [5.5, 5.75, 5.75]
    for row in tbl.rows:
        for j, w in enumerate(widths):
            _set_col_width(row.cells[j], w)

    # Título
    tbl.rows[0].cells[0].merge(tbl.rows[0].cells[2])
    _cell_text(tbl.rows[0].cells[0], "PROMEDIOS GLOBALES", bold=True, size=9)
    _set_cell_bg(tbl.rows[0].cells[0], C_GREEN_TITLE)

    # Encabezado
    hdr = tbl.rows[1]
    for j, (txt, fill) in enumerate([
        ("Indicador",  C_GREEN_HEADER),
        ("FUENTE (A–E)", C_GREEN_HEADER),
        ("FONDO (I–V)",  C_BLUE_HEADER),
    ]):
        _cell_text(hdr.cells[j], txt, bold=True, size=9)
        _set_cell_bg(hdr.cells[j], fill)

    indicadores = [
        ("N₅₀ (dB)", "N50"), ("σ (dB)", "sigma"),
        ("N₁₀ (dB)", "N10"), ("(Neq)eq (dB)", "Neq"),
    ]
    for i, (label, key) in enumerate(indicadores):
        row = tbl.rows[2 + i]
        fill = C_WHITE if i % 2 == 0 else C_GREY_ROW
        _cell_text(row.cells[0], label, bold=True, size=9, align="left")
        _set_cell_bg(row.cells[0], C_GREEN_HEADER)
        _cell_text(row.cells[1], f"{prom['fuente'][key]:.2f}", size=9)
        _set_cell_bg(row.cells[1], fill)
        _cell_text(row.cells[2], f"{prom['fondo'][key]:.2f}", size=9)
        _set_cell_bg(row.cells[2], fill)

    _style_all_cells(tbl)


def _tabla_correcciones(doc, corr, prom):
    """Tabla de parámetros de corrección con valores."""
    tbl = doc.add_table(rows=1 + 4, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [7.0, 4.0, 6.0]
    for row in tbl.rows:
        for j, w in enumerate(widths):
            _set_col_width(row.cells[j], w)

    hdr = tbl.rows[0]
    for j, (txt, fill) in enumerate([
        ("Parámetro", C_GREEN_HEADER),
        ("Valor",     C_GREEN_HEADER),
        ("Unidad",    C_GREEN_HEADER),
    ]):
        _cell_text(hdr.cells[j], txt, bold=True, size=9)
        _set_cell_bg(hdr.cells[j], fill)

    cf_val = f"{corr['Cf']:.2f}" if corr["Cf_aplica"] else "No Aplica"
    n50c_val = f"{corr['N50_corr']:.2f}" if corr["Cf_aplica"] else "No Aplica"

    filas = [
        ("Ce = 0.9023 × σ_prom", f"{corr['Ce']:.2f}", "dB"),
        ("Δ₅₀ = N₅₀_fuente − N₅₀_fondo",
         f"{abs(corr['delta50']):.2f}", "dB"),
        ("Cf",  cf_val,  "dB"),
        ("N'₅₀", n50c_val, "dB"),
    ]
    for i, (p_txt, v_txt, u_txt) in enumerate(filas):
        row = tbl.rows[1 + i]
        fill = C_WHITE if i % 2 == 0 else C_GREY_ROW
        _cell_text(row.cells[0], p_txt, size=9, align="left")
        _set_cell_bg(row.cells[0], fill)
        _cell_text(row.cells[1], v_txt, bold=True, size=9)
        _set_cell_bg(row.cells[1], fill)
        _cell_text(row.cells[2], u_txt, size=9)
        _set_cell_bg(row.cells[2], fill)

    _style_all_cells(tbl)


def _tabla_resultado_final(doc, res, limite):
    """Tabla de resultado final con semáforo CUMPLE / EXCEDE."""
    excede    = res["Nff_corr"] > limite
    fill_res  = C_RED_FAIL if excede else C_GREEN_PASS
    color_txt = C_RESULT_TEXT_FAIL if excede else C_RESULT_TEXT_OK
    verdict   = (f"EXCEDE el límite ({limite:.1f} dB)"
                 if excede else
                 f"CUMPLE el límite ({limite:.1f} dB)")

    tbl = doc.add_table(rows=1 + 2 + 1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [7.0, 4.0, 6.0]
    for row in tbl.rows:
        for j, w in enumerate(widths):
            _set_col_width(row.cells[j], w)

    # Título
    tbl.rows[0].cells[0].merge(tbl.rows[0].cells[2])
    _cell_text(tbl.rows[0].cells[0],
               "NIVEL DE FUENTE FIJA — RESULTADO FINAL",
               bold=True, size=10)
    _set_cell_bg(tbl.rows[0].cells[0], C_GREEN_TITLE)

    # Nff
    r1 = tbl.rows[1]
    _cell_text(r1.cells[0], "N_ff =",    size=10, align="left")
    _cell_text(r1.cells[1], f"{res['Nff']:.2f}", bold=True, size=11)
    _cell_text(r1.cells[2], "dB",        size=10)
    for cell in r1.cells:
        _set_cell_bg(cell, C_WHITE)

    # (N')ff  — resultado con color
    r2 = tbl.rows[2]
    _cell_text(r2.cells[0], "(N')ff =",  bold=True, size=11, align="left")
    _cell_text(r2.cells[1], f"{res['Nff_corr']:.2f}",
               bold=True, size=13, color=color_txt)
    _cell_text(r2.cells[2], "dB", bold=True, size=11)
    for cell in r2.cells:
        _set_cell_bg(cell, fill_res)

    # Veredicto
    r3 = tbl.rows[3]
    r3.cells[0].merge(r3.cells[2])
    _cell_text(r3.cells[0], verdict, bold=True, size=11, color=color_txt)
    _set_cell_bg(r3.cells[0], fill_res)

    _style_all_cells(tbl)


# ╔══════════════════════════════════════════════════════════════════╗
# ║  FUNCIÓN PRINCIPAL                                              ║
# ╚══════════════════════════════════════════════════════════════════╝

def generar_word(resultado: dict) -> bytes:
    """
    Genera la Memoria de Cálculo NOM-081 en formato Word.

    Parámetros
    ----------
    resultado : dict
        Estructura con claves:
        metadata, fuente_data, fondo_data, fuente_stats, fondo_stats,
        promedios, correcciones, resultado

    Retorna
    -------
    bytes  →  contenido del .docx listo para descarga
    """
    meta  = resultado["metadata"]
    fd    = resultado["fuente_data"]
    bd    = resultado["fondo_data"]
    fs    = resultado["fuente_stats"]
    bs    = resultado["fondo_stats"]
    prom  = resultado["promedios"]
    corr  = resultado["correcciones"]
    res   = resultado["resultado"]
    limite = float(meta.get("limite", 68.0))

    per_f = list(fd.keys())   # A B C D E
    per_b = list(bd.keys())   # I II III IV V

    doc = Document()
    _setup_page(doc)
    _add_footer(doc, meta)

    # ── Configurar estilo Normal ──────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════
    # TÍTULO PRINCIPAL
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4); pf.space_after = Pt(6)
    # Sombreado del párrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_hex(C_GREEN_TITLE))
    # shd must come before spacing, ind, jc, tabs, rPr in pPr
    inserted = False
    for anchor_tag in [qn("w:spacing"), qn("w:ind"),
                       qn("w:tabs"), qn("w:jc"), qn("w:rPr")]:
        anchor = pPr.find(anchor_tag)
        if anchor is not None:
            anchor.addprevious(shd)
            inserted = True
            break
    if not inserted:
        pPr.append(shd)
    run = p.add_run("CÁLCULOS PARA ZONA CRITICA 1")
    run.bold = True; run.font.name = FONT_NAME; run.font.size = Pt(12)

    # ── Ficha de identificación ───────────────────────────────────
    _add_para(doc, "", space_before=4, space_after=2)
    _tabla_encabezado(doc, meta)
    _add_para(doc, "", space_before=6, space_after=2)

    # ══════════════════════════════════════════════════════════════
    # DATOS DE CAMPO
    # ══════════════════════════════════════════════════════════════
    _tabla_datos_campo(doc, "FUENTE", per_f, fd, fs)
    _add_para(doc, "", space_before=8, space_after=2)
    _tabla_datos_campo(doc, "FONDO",  per_b, bd, bs)
    _add_para(doc, "", space_before=8, space_after=2)

    # ══════════════════════════════════════════════════════════════
    # A. FÓRMULAS APLICABLES
    # ══════════════════════════════════════════════════════════════
    _add_para(doc, "A.  Fórmulas aplicables",
              bold=True, size=11, space_before=10, space_after=4)

    _add_para(doc,
              "Para cada periodo de 35 lecturas (Nᵢ) en dB se calculan:",
              size=10, space_before=2, space_after=2)

    formulas_a = [
        ("N₅₀ = (1/n) × Σ Nᵢ",
         "N₅₀: Nivel sonoro medio aritmético de las 35 lecturas."),
        ("σ = √[ Σ(Nᵢ − N₅₀)² / (n−1) ]",
         "σ: Desviación estándar muestral (ddof = 1, n = 35)."),
        ("N₁₀ = N₅₀ + 1.2817 × σ",
         "N₁₀: Nivel superado el 10 % del tiempo (percentil 90). "
         "El factor 1.2817 es el valor z de la normal estándar para P = 90 %."),
        ("Neq = 10 × log₁₀[ (1/n) × Σ 10^(Nᵢ/10) ]",
         "Neq: Nivel de presión sonora equivalente (promedio energético)."),
    ]
    for formula, expl in formulas_a:
        _add_formula(doc, formula)
        _add_indent(doc, expl, size=9)

    _add_para(doc, "Promedios globales de los 5 periodos:",
              bold=True, size=10, space_before=6, space_after=2)

    formulas_prom = [
        ("N₅₀_prom = (1/5) × Σ N₅₀ᵢ",
         "Media aritmética de los cinco N₅₀."),
        ("σ_prom = (1/5) × Σ σᵢ",
         "Media aritmética de las cinco desviaciones estándar."),
        ("N₁₀_prom = N₅₀_prom + 1.2817 × σ_prom",
         "Nivel percentil 90 global."),
        ("(Neq)eq = 10 × log₁₀[ (1/5) × Σ 10^(Neqᵢ/10) ]",
         "(Neq)eq: Promedio energético de los Neq de los 5 periodos."),
    ]
    for formula, expl in formulas_prom:
        _add_formula(doc, formula)
        _add_indent(doc, expl, size=9)

    # ══════════════════════════════════════════════════════════════
    # B. SUSTITUCIÓN
    # ══════════════════════════════════════════════════════════════
    _add_para(doc, "B.  Sustitución",
              bold=True, size=11, space_before=10, space_after=4)
    _add_para(doc,
              "Aplicando las fórmulas a los datos medidos en campo:",
              size=10, space_before=2, space_after=4)

    _tabla_resultados_periodo(doc, "FUENTE", per_f, fs, C_GREEN_TITLE)
    _add_para(doc, "", space_before=6, space_after=2)
    _tabla_resultados_periodo(doc, "FONDO",  per_b, bs, C_BLUE_HEADER)
    _add_para(doc, "", space_before=6, space_after=2)
    _tabla_promedios(doc, prom)
    _add_para(doc, "", space_before=6, space_after=2)

    # ══════════════════════════════════════════════════════════════
    # C. CÁLCULO DE CORRECCIONES
    # ══════════════════════════════════════════════════════════════
    _add_para(doc, "C.  Cálculo de correcciones",
              bold=True, size=11, space_before=10, space_after=4)
    _add_para(doc,
              "Una vez obtenidos los promedios, se aplican las correcciones de la norma:",
              size=10, space_before=2, space_after=4)

    # Explicación paso a paso
    _add_indent(doc,
        f"Ce = 0.9023 × σ_prom(fuente)  =  "
        f"0.9023 × {prom['fuente']['sigma']:.2f}  =  "
        f"{corr['Ce']:.2f} dB",
        size=10)
    _add_indent(doc,
        "Ce: Corrección por presencia de valores extremos "
        "en la distribución de niveles.", size=9)

    _add_indent(doc,
        f"Δ₅₀ = N₅₀_fuente − N₅₀_fondo  =  "
        f"{prom['fuente']['N50']:.2f} − {prom['fondo']['N50']:.2f}  =  "
        f"{corr['delta50']:.2f} dB",
        size=10)
    _add_indent(doc,
        "Δ₅₀: Diferencia entre el nivel medio de la fuente y el fondo. "
        "Si Δ₅₀ ≥ 0.75 dB aplica la corrección por ruido de fondo (Cf).",
        size=9)

    _add_indent(doc,
        f"N'₅₀ = N₅₀_prom + Ce  =  "
        f"{prom['fuente']['N50']:.2f} + {corr['Ce']:.2f}  =  "
        f"{corr['N50_corr']:.2f} dB",
        size=10)
    _add_indent(doc,
        "N'₅₀: Nivel N₅₀ corregido por los valores extremos de la muestra.",
        size=9)

    _add_para(doc, "", space_before=4, space_after=2)
    _tabla_correcciones(doc, corr, prom)
    _add_para(doc, "", space_before=4, space_after=2)

    # Explicación Cf
    if not corr["Cf_aplica"]:
        _add_indent(doc,
            f"Como Δ₅₀ = {corr['delta50']:.2f} dB < 0.75 dB, "
            f"la corrección por ruido de fondo NO APLICA.",
            size=10, bold=True)
    else:
        _add_indent(doc,
            f"Como Δ₅₀ = {corr['delta50']:.2f} dB ≥ 0.75 dB, aplica Cf:",
            size=10)
        _add_formula(doc,
            f"Cf = −(Δ₅₀+9) + 3√(4·Δ₅₀−3)  =  {corr['Cf']:.2f} dB")

    # ══════════════════════════════════════════════════════════════
    # D. DETERMINACIÓN DEL Nff
    # ══════════════════════════════════════════════════════════════
    _add_para(doc, "D.  Determinación del nivel de fuente fija (Nff)",
              bold=True, size=11, space_before=10, space_after=4)
    _add_para(doc,
              "El N_ff es el mayor valor entre N'₅₀ y (Neq)eq:",
              size=10, space_before=2, space_after=2)

    _add_formula(doc,
        f"Nff = max( N'₅₀ , (Neq)eq )  =  "
        f"max( {corr['N50_corr']:.2f} , {prom['fuente']['Neq']:.2f} )  =  "
        f"{res['Nff']:.2f} dB")

    _add_para(doc, "", space_before=4, space_after=2)

    if not corr["Cf_aplica"]:
        _add_indent(doc,
            f"Dado que Cf no aplica (Δ₅₀ = {corr['delta50']:.2f} dB < 0.75 dB):",
            size=10)
        _add_formula(doc,
            f"(N')ff = Nff = {res['Nff_corr']:.2f} dB")
    else:
        _add_indent(doc,
            "Aplicando la corrección por ruido de fondo:", size=10)
        _add_formula(doc,
            f"(N')ff = Nff + Cf  =  "
            f"{res['Nff']:.2f} + {corr['Cf']:.2f}  =  "
            f"{res['Nff_corr']:.2f} dB")

    _add_para(doc, "", space_before=6, space_after=2)
    _tabla_resultado_final(doc, res, limite)

    # ── Serializar a bytes ────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    raw = buf.read()

    # Post-proceso: corregir w:zoom que python-docx genera sin w:percent
    import zipfile, re as _re
    in_zip  = zipfile.ZipFile(io.BytesIO(raw))
    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as out_zip:
        for item in in_zip.infolist():
            data = in_zip.read(item.filename)
            if item.filename == "word/settings.xml":
                # Añadir w:percent="100" si falta
                data = data.replace(
                    b'<w:zoom w:val=',
                    b'<w:zoom w:percent="100" w:val='
                )
            out_zip.writestr(item, data)
    out_buf.seek(0)
    return out_buf.read()
